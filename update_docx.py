import os
import sys
import subprocess

# Explicitly uninstall legacy 'docx' if it got installed by mistake, and install 'python-docx'
def setup_packages():
    try:
        # Uninstall legacy docx if present to prevent namespace conflicts
        subprocess.check_call([sys.executable, "-m", "pip", "uninstall", "-y", "docx"])
    except Exception:
        pass
    
    try:
        # Install modern python-docx
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    except Exception as e:
        print(f"Error installing python-docx: {e}")
        sys.exit(1)

setup_packages()

from docx import Document
from docx.shared import Pt, RGBColor

def add_custom_heading(doc, text, size=14, color=None, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def main():
    docx_path = "Project.docx"
    if not os.path.exists(docx_path):
        print(f"Creating a new Document since {docx_path} was not found.")
        doc = Document()
    else:
        print(f"Opening existing {docx_path}...")
        doc = Document(docx_path)

    # Purple color accent
    purple_accent = RGBColor(138, 43, 226) # Purple
    dark_gray = RGBColor(50, 50, 50)

    # Add a page break if the document is not empty
    if len(doc.paragraphs) > 0:
        doc.add_page_break()

    # Add Heading
    add_custom_heading(doc, 'Project Technical Details & Architecture Reference', size=18, color=purple_accent, bold=True)

    # Add intro paragraph
    p = doc.add_paragraph('This document outlines the technical architecture, database schema, API endpoints, and core AI integrations of the Eventify Event Management System.')
    p.runs[0].font.name = 'Arial'

    # Section 1: Database Tables
    add_custom_heading(doc, '1. Database Schema', size=15, color=purple_accent, bold=True)
    
    table_intro = doc.add_paragraph('The application utilizes a MySQL database. Below are the key tables used:')
    table_intro.runs[0].font.name = 'Arial'
    
    # Tables details
    tables = [
        ("users", "Stores user profiles and roles (regular users vs managers).", 
         "id (INT, PK, Auto-Increment)\nname (VARCHAR)\nemail (VARCHAR, Unique)\npassword (VARCHAR, Hashed)\nrole (VARCHAR, 'user' or 'manager')\ncreated_at (TIMESTAMP)"),
        ("events", "Stores event details, venues, ticket pricing, and manager ownership.", 
         "id (INT, PK, Auto-Increment)\nname (VARCHAR)\ndescription (TEXT)\ndate (VARCHAR)\ntime (VARCHAR)\nvenue (VARCHAR)\nticket_price (INT)\ncreated_by (INT, FK to users)\ncreated_at (TIMESTAMP)"),
        ("bookings", "Stores event slot bookings registered by users.", 
         "id (INT, PK, Auto-Increment)\nname (VARCHAR)\nemail (VARCHAR)\nphone (VARCHAR)\nevent_id (INT, FK to events)\ncreated_at (TIMESTAMP)"),
        ("payments", "Stores payment logs and transaction amounts for event tickets.", 
         "id (INT, PK, Auto-Increment)\nevent_id (INT)\namount (DECIMAL)\ncreated_at (TIMESTAMP)")
    ]

    for name, desc, columns in tables:
        add_custom_heading(doc, f"Table: {name}", size=13, color=dark_gray, bold=True)
        p_desc = doc.add_paragraph(desc)
        p_desc.runs[0].font.name = 'Arial'
        p_cols = doc.add_paragraph()
        run_cols = p_cols.add_run(f"Columns:\n{columns}")
        run_cols.font.name = 'Courier New'
        run_cols.font.size = Pt(10)

    # Section 2: API Endpoints
    add_custom_heading(doc, '2. Backend API Endpoints', size=15, color=purple_accent, bold=True)
    
    apis = [
        ("Authentication", [
            ("POST /api/auth/register", "Registers a new user (role defaults to 'user')."),
            ("POST /api/auth/login", "Authenticates a user and returns a JWT token.")
        ]),
        ("Events", [
            ("GET /api/events/events", "Fetches all events. If logged in as a manager, only returns events they created."),
            ("POST /api/events/events", "Creates a new event. (Requires manager token)"),
            ("PUT /api/events/events/:id/description", "Updates the description of an event. (Requires manager token)"),
            ("DELETE /api/events/events/:id", "Deletes an event. (Requires manager token)")
        ]),
        ("Bookings", [
            ("GET /api/bookings/bookings", "Fetches slot bookings. Regular users see their bookings; managers see all bookings."),
            ("POST /api/bookings/bookings", "Creates a new slot booking. (Requires authenticated token)")
        ]),
        ("Payments", [
            ("POST /api/payment/create-order", "Creates a new Razorpay order. Falls back to generating a mock order if Razorpay is offline."),
            ("POST /api/payment/save", "Records a transaction in the payments database table."),
            ("GET /api/payment/orders", "Fetches transaction logs for display in the payments dashboard tab.")
        ]),
        ("AI Integrations", [
            ("POST /api/ai/generate-description", "Generates a descriptive catchphrase for events using title, date, time, and venue context."),
            ("POST /api/ai/chat", "Interactive endpoint for the Eventify chatbot. It reads active events from the SQL database and passes them to Gemini for context.")
        ])
    ]

    for category, endpoints in apis:
        add_custom_heading(doc, category, size=13, color=dark_gray, bold=True)
        for route, desc in endpoints:
            p_route = doc.add_paragraph()
            r_route = p_route.add_run(f"{route}\n")
            r_route.font.bold = True
            r_route.font.name = 'Courier New'
            r_route.font.size = Pt(10)
            r_desc = p_route.add_run(desc)
            r_desc.font.name = 'Arial'

    # Section 3: AI Core Features
    add_custom_heading(doc, '3. Core AI Integration & Features', size=15, color=purple_accent, bold=True)
    
    add_custom_heading(doc, 'AI Event Description Generator', size=13, color=dark_gray, bold=True)
    p_desc_gen = doc.add_paragraph(
        "Built using Google's Gemini 2.5 Flash API, this feature allows managers to generate catchy, "
        "personalized event descriptions instantly. It accepts metadata details (name, date, time, and venue) "
        "and formulates an engaging paragraph structure, streamlining the event registration process."
    )
    p_desc_gen.runs[0].font.name = 'Arial'

    add_custom_heading(doc, 'Interactive AI Event Chatbot Assistant', size=13, color=dark_gray, bold=True)
    p_chatbot = doc.add_paragraph(
        "A premium glassmorphic floating chatbot widget integrated directly into the user dashboard. "
        "The backend dynamically queries all events in the MySQL database and feeds this context "
        "into the Gemini API system instructions. Users can ask questions in natural language, "
        "request event recommendations, search for free/paid events, or query ticket prices. "
        "It supports message history preservation and auto-scrolling."
    )
    p_chatbot.runs[0].font.name = 'Arial'

    # Save
    doc.save(docx_path)
    print(f"Successfully updated {docx_path}!")

if __name__ == "__main__":
    main()
